#!/usr/bin/env python3
"""
MVP Document Processor - Enhanced Flask Web Application
Complete web interface with user inputs, comprehensive analysis, and disclaimer bookmarks
"""

from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
import os
import tempfile
import yaml
import re
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import textstat
from collections import defaultdict
import io
from pathlib import Path
import traceback

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'mvp-processor-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Allowed file extensions
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class MVPDocumentProcessor:
    """
    Enhanced MVP Document Processor with user inputs, comprehensive analysis, and disclaimer bookmarks
    """
    
    def __init__(self):
        self.corrections_made = []
        self.statistics = {
            'total_corrections': 0,
            'rules_applied': defaultdict(int),
            'word_count': 0,
            'sentence_count': 0,
            'paragraph_count': 0,
            'reading_level': 0
        }
        self.detailed_corrections = []
        self.corrections_by_category = defaultdict(int)
        self.rules = self._load_mvp_rules()
        self.medicare_checks = []
        self.keyword_analysis = {}
        self.user_config = {}
    
    def _load_mvp_rules(self):
        """Load all MVP corporate rules"""
        return {
            'time_formatting_rules': [
                {
                    'category': 'remove_all_unnecessary_minutes',
                    'find': r'\b(\d{1,2}):00\b',
                    'replace': r'\1',
                    'case_sensitive': False,
                    'description': "Remove all instances of :00 minutes",
                    'enabled': True
                },
                {
                    'category': 'am_all_variations_lowercase',
                    'find': r'\b(\d{1,2}(?::\d{2})?)\s*([Aa]\.?[Mm]\.?)\b',
                    'replace': r'\1 am',
                    'case_sensitive': False,
                    'description': "Convert all AM variations to lowercase am",
                    'enabled': True
                },
                {
                    'category': 'pm_all_variations_lowercase',
                    'find': r'\b(\d{1,2}(?::\d{2})?)\s*([Pp]\.?[Mm]\.?)\b',
                    'replace': r'\1 pm',
                    'case_sensitive': False,
                    'description': "Convert all PM variations to lowercase pm",
                    'enabled': True
                },
                {
                    'category': 'time_range_en_dash',
                    'find': r'\b(\d{1,2}(?::\d{2})?\s*(?:am|pm))\s*[-–—]\s*(\d{1,2}(?::\d{2})?\s*(?:am|pm))\b',
                    'replace': r'\1–\2',
                    'case_sensitive': False,
                    'description': "Use en dash with no spaces for time ranges",
                    'enabled': True
                },
                {
                    'category': 'space_before_am_pm',
                    'find': r'\b(\d{1,2}(?::\d{2})?)(am|pm)\b',
                    'replace': r'\1 \2',
                    'case_sensitive': False,
                    'description': "Ensure space between number and am/pm",
                    'enabled': True
                }
            ],
            'number_formatting_rules': [
                {
                    'category': 'spell_out_small_numbers',
                    'find': r'\b(?<![\d\-/])(?<!January\s)(?<!February\s)(?<!March\s)(?<!April\s)(?<!May\s)(?<!June\s)(?<!July\s)(?<!August\s)(?<!September\s)(?<!October\s)(?<!November\s)(?<!December\s)([1-9])\b(?!\s*(?:[AaPp]\.?[Mm]\.?|am|pm|:\d|%|\.|,\d{3}|-star|th|nd|rd|st)\b)(?![\d\-/])',
                    'replace': 'NUMBER_WORD_\\1',
                    'case_sensitive': False,
                    'description': "Spell out numbers 1-9 (with exclusions)",
                    'enabled': True
                },
                {
                    'category': 'comma_in_large_numbers',
                    'find': r'\b(?<!extension\s)(?<!ext\.\s)(?<![-.])\b(?!20[2-5]\d\b)(\d{1,3})(\d{3})\b(?![\d/])',
                    'replace': r'\1,\2',
                    'case_sensitive': False,
                    'description': "Add commas to numbers 1,000+ (excluding years 2020-2050 and phone numbers)",
                    'enabled': True
                }
            ],
            'brand_trademark_rules': [
                {
                    'category': 'gia_registration_mark',
                    'find': r'\bGia(?!®)\b',
                    'replace': 'Gia®',
                    'case_sensitive': True,
                    'description': "Add registration mark to first instance of Gia",
                    'first_instance_only': True,
                    'enabled': True
                }
            ],
            'mvp_terminology_rules': [
                {
                    'category': 'mvp_health_plans_to_mvp_health_care_plans',
                    'find': r'\bMVP health plans\b',
                    'replace': 'MVP Health Care plans',
                    'case_sensitive': False,
                    'description': "Use full company name for plans",
                    'enabled': True
                },
                {
                    'category': 'telehealth_to_virtual_care',
                    'find': r'\btelehealth\b',
                    'replace': 'virtual care',
                    'case_sensitive': False,
                    'description': "Use virtual care for member communications",
                    'enabled': True
                },
                {
                    'category': 'healthcare_terminology',
                    'find': r'\bhealthcare\b',
                    'replace': 'health care',
                    'case_sensitive': False,
                    'description': "Always use 'health care' (two words)",
                    'enabled': True
                },
                {
                    'category': 'login_to_signin',
                    'find': r'\blogin\b',
                    'replace': 'sign in',
                    'case_sensitive': False,
                    'description': "Replace 'login' with 'sign in'",
                    'enabled': True
                },
                {
                    'category': 'log_in_to_sign_in',
                    'find': r'\blog in\b(?!\s+to)',
                    'replace': 'sign in',
                    'case_sensitive': False,
                    'description': "Replace 'log in' with 'sign in'",
                    'enabled': True
                },
                {
                    'category': 'preventative_to_preventive',
                    'find': r'\bpreventative\b',
                    'replace': 'preventive',
                    'case_sensitive': False,
                    'description': "Use 'preventive' instead of 'preventative'",
                    'enabled': True
                }
            ],
            'punctuation_rules': [
                {
                    'category': 'ampersand_replacement',
                    'find': r'\s&\s',
                    'replace': ' and ',
                    'case_sensitive': False,
                    'description': "Replace ampersands with 'and'",
                    'enabled': True
                },
                {
                    'category': 'double_spaces',
                    'find': r'  +',
                    'replace': ' ',
                    'case_sensitive': False,
                    'description': "Remove multiple spaces",
                    'enabled': True
                }
            ],
            'state_abbreviation_rules': [
                {
                    'category': 'remove_periods_from_states',
                    'find': r'\b([NnVvCc])\.([YyTt])\.\b',
                    'replace': r'\1\2',
                    'case_sensitive': False,
                    'description': "Remove periods from state abbreviations (NY, VT, CT only)",
                    'enabled': True
                },
                {
                    'category': 'capitalize_state_abbreviations',
                    'find': r'\b(ny|vt|ct)\b',
                    'replace': lambda m: m.group(1).upper(),
                    'case_sensitive': False,
                    'description': "Capitalize state abbreviations (NY, VT, CT only)",
                    'enabled': True,
                    'is_function': True
                }
            ],
            'medicare_rules': [
                {
                    'category': 'add_tty_to_phone_numbers',
                    'find': r'\b(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})\b(?!\s*\(TTY 711\))',
                    'replace': r'\1 (TTY 711)',
                    'case_sensitive': False,
                    'description': "Add (TTY 711) after phone numbers for Medicare pages",
                    'enabled': False  # Will be enabled when is_medicare_page = True
                }
            ]
        }
    
    def _extract_angle_bracket_content(self, text):
        """Extract and preserve content in angle brackets"""
        angle_bracket_pattern = r'<[^>]*>'
        brackets = []
        bracket_placeholders = {}
        
        def extract_bracket(match):
            placeholder = f"ANGLE_BRACKET_PLACEHOLDER_{len(brackets)}"
            brackets.append(match.group(0))
            bracket_placeholders[placeholder] = match.group(0)
            return placeholder
        
        text_without_brackets = re.sub(angle_bracket_pattern, extract_bracket, text)
        return text_without_brackets, bracket_placeholders
    
    def _restore_angle_bracket_content(self, text, bracket_placeholders):
        """Restore angle bracket content from placeholders"""
        for placeholder, original_bracket in bracket_placeholders.items():
            text = text.replace(placeholder, original_bracket)
        return text
    
    def _find_bookmark_range(self, doc, start_bookmark="start_page_copy", end_bookmark="end_page_copy"):
        """Find paragraph range between text markers (not Word bookmarks)"""
        start_para = None
        end_para = None
        
        # Search for text strings in paragraph content
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip().lower()
            
            # Check if paragraph contains start marker
            if start_bookmark.lower() in para_text:
                start_para = i
                print(f"📍 Found start marker '{start_bookmark}' at paragraph {i}")
            
            # Check if paragraph contains end marker
            if end_bookmark.lower() in para_text:
                end_para = i
                print(f"📍 Found end marker '{end_bookmark}' at paragraph {i}")
                break  # Stop searching after finding end marker
        
        return start_para, end_para

    def _find_disclaimer_range(self, doc, start_disclaimer="start_disclaimer", end_disclaimer="end_disclaimer"):
        """Find paragraph range between disclaimer markers"""
        start_para = None
        end_para = None
        
        # Search for disclaimer markers in paragraph content
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip().lower()
            
            # Check if paragraph contains start disclaimer marker
            if start_disclaimer.lower() in para_text:
                start_para = i
                print(f"📍 Found disclaimer start marker '{start_disclaimer}' at paragraph {i}")
            
            # Check if paragraph contains end disclaimer marker
            if end_disclaimer.lower() in para_text:
                end_para = i
                print(f"📍 Found disclaimer end marker '{end_disclaimer}' at paragraph {i}")
                break  # Stop searching after finding end marker
        
        return start_para, end_para
    
    def _check_medicare_compliance(self, doc):
        """Check Medicare-specific compliance requirements with enhanced disclaimer checking"""
        medicare_issues = []
        
        # Get full document text for general checks
        full_text = ' '.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Check for CMS code specifically in disclaimer section if Medicare page
        if self.user_config.get('is_medicare_page'):
            # Find disclaimer section
            disclaimer_start, disclaimer_end = self._find_disclaimer_range(doc)
            
            if disclaimer_start is not None and disclaimer_end is not None:
                # Get text only from disclaimer section
                disclaimer_paragraphs = doc.paragraphs[disclaimer_start:disclaimer_end + 1]
                disclaimer_text = ' '.join([p.text for p in disclaimer_paragraphs if p.text.strip()])
                
                print(f"📍 Checking CMS code in disclaimer section (paragraphs {disclaimer_start}-{disclaimer_end})")
                
                # Check for CMS code in disclaimer section
                cms_code_pattern = r'\bY[A-Z0-9]*_[A-Z0-9]*_[A-Z0-9]*\b'
                cms_codes = re.findall(cms_code_pattern, disclaimer_text)
                
                if not cms_codes:
                    medicare_issues.append({
                        'type': 'missing_cms_code_in_disclaimer',
                        'description': 'Missing CMS code in disclaimer section (should start with Y and contain two underscores)',
                        'severity': 'high',
                        'location': f'Disclaimer section (paragraphs {disclaimer_start}-{disclaimer_end})'
                    })
                else:
                    medicare_issues.append({
                        'type': 'cms_code_found_in_disclaimer',
                        'description': f'CMS code found in disclaimer: {cms_codes[0]}',
                        'severity': 'info',
                        'location': f'Disclaimer section (paragraphs {disclaimer_start}-{disclaimer_end})'
                    })
            else:
                # If no disclaimer section found, check entire document
                print(f"📍 No disclaimer section found, checking entire document for CMS code")
                cms_code_pattern = r'\bY[A-Z0-9]*_[A-Z0-9]*_[A-Z0-9]*\b'
                cms_codes = re.findall(cms_code_pattern, full_text)
                
                if not cms_codes:
                    medicare_issues.append({
                        'type': 'missing_cms_code',
                        'description': 'Missing CMS code (should start with Y and contain two underscores). Consider adding disclaimer section with start_disclaimer and end_disclaimer markers.',
                        'severity': 'high',
                        'location': 'Entire document'
                    })
                else:
                    medicare_issues.append({
                        'type': 'cms_code_found',
                        'description': f'CMS code found: {cms_codes[0]} (Note: Consider placing in disclaimer section)',
                        'severity': 'info',
                        'location': 'Document body'
                    })
        
        # Check for phone numbers without TTY (in main content area)
        # Find main content area
        content_start, content_end = self._find_bookmark_range(doc)
        
        if content_start is not None and content_end is not None:
            # Check only main content area for phone numbers
            content_paragraphs = doc.paragraphs[content_start:content_end + 1]
            content_text = ' '.join([p.text for p in content_paragraphs if p.text.strip()])
            
            phone_pattern = r'\b(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})\b(?!\s*\(TTY 711\))'
            phone_matches = re.findall(phone_pattern, content_text)
            
            if phone_matches:
                medicare_issues.append({
                    'type': 'phone_missing_tty_in_content',
                    'description': f'Found {len(phone_matches)} phone number(s) without TTY 711 in main content area',
                    'severity': 'medium',
                    'phone_numbers': phone_matches,
                    'location': f'Main content area (paragraphs {content_start}-{content_end})'
                })
        else:
            # Check entire document if no content markers found
            phone_pattern = r'\b(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})\b(?!\s*\(TTY 711\))'
            phone_matches = re.findall(phone_pattern, full_text)
            
            if phone_matches:
                medicare_issues.append({
                    'type': 'phone_missing_tty',
                    'description': f'Found {len(phone_matches)} phone number(s) without TTY 711',
                    'severity': 'medium',
                    'phone_numbers': phone_matches,
                    'location': 'Entire document'
                })
        
        return medicare_issues
    
    def _analyze_keywords(self, text, keywords):
        """Analyze keyword frequency in text"""
        keyword_analysis = {}
        text_lower = text.lower()
        
        for keyword in keywords:
            if keyword.strip():
                keyword_lower = keyword.strip().lower()
                # Count whole word matches
                pattern = r'\b' + re.escape(keyword_lower) + r'\b'
                matches = len(re.findall(pattern, text_lower))
                keyword_analysis[keyword.strip()] = matches
        
        return keyword_analysis
    
    def _create_analysis_report(self, doc, results):
        """Create comprehensive analysis report and append to document"""
        # Add page break
        doc.add_page_break()
        
        # Title
        title = doc.add_heading('Document Processing Analysis Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_run = timestamp_para.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        timestamp_run.italic = True
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        
        # Configuration used
        config_para = doc.add_paragraph()
        config_para.add_run('Processing Configuration: ').bold = True
        config_text = f"Target: {self.user_config.get('target_word_count', 'Not specified')} words, "
        config_text += f"Reading Level: {self.user_config.get('target_reading_level', 'Not specified')}, "
        config_text += f"Medicare Page: {'Yes' if self.user_config.get('is_medicare_page') else 'No'}"
        config_para.add_run(config_text)
        
        # Document structure analysis
        content_start, content_end = self._find_bookmark_range(doc)
        disclaimer_start, disclaimer_end = self._find_disclaimer_range(doc)
        
        structure_para = doc.add_paragraph()
        structure_para.add_run('Document Structure: ').bold = True
        structure_text = ""
        if content_start is not None and content_end is not None:
            structure_text += f"Main content area found (paragraphs {content_start}-{content_end}). "
        else:
            structure_text += "No content markers found, processed entire document. "
        
        if disclaimer_start is not None and disclaimer_end is not None:
            structure_text += f"Disclaimer section found (paragraphs {disclaimer_start}-{disclaimer_end})."
        else:
            structure_text += "No disclaimer section found."
        
        structure_para.add_run(structure_text)
        
        # Statistics Summary
        stats = results['document_statistics']
        summary_para = doc.add_paragraph()
        summary_para.add_run('Document Statistics: ').bold = True
        summary_text = (f"{stats['word_count']:,} words • "
                       f"{stats['sentence_count']} sentences • "
                       f"Grade {stats['reading_level']:.1f} reading level • "
                       f"{results['total_corrections']} corrections applied")
        summary_para.add_run(summary_text)
        
        # Word Count Analysis
        doc.add_heading('Word Count Analysis', level=2)
        
        target_words = self.user_config.get('target_word_count')
        actual_words = stats['word_count']
        
        if target_words:
            word_diff = actual_words - target_words
            percentage = (actual_words / target_words) * 100
            
            word_analysis_para = doc.add_paragraph()
            word_analysis_para.add_run('Target Word Count: ').bold = True
            word_analysis_para.add_run(f"{target_words:,} words")
            
            actual_para = doc.add_paragraph()
            actual_para.add_run('Actual Word Count: ').bold = True
            actual_para.add_run(f"{actual_words:,} words ({percentage:.1f}% of target)")
            
            diff_para = doc.add_paragraph()
            diff_para.add_run('Difference: ').bold = True
            if word_diff > 0:
                diff_para.add_run(f"+{word_diff:,} words over target")
            elif word_diff < 0:
                diff_para.add_run(f"{word_diff:,} words under target")
            else:
                diff_para.add_run("Exactly on target")
        else:
            no_target_para = doc.add_paragraph()
            no_target_para.add_run('No target word count specified. ')
            no_target_para.add_run(f'Current word count: {actual_words:,} words').bold = True
        
        # Reading Level Analysis
        doc.add_heading('Reading Level Analysis', level=2)
        
        target_level = self.user_config.get('target_reading_level')
        actual_level = stats['reading_level']
        
        if target_level:
            level_diff = actual_level - target_level
            
            target_level_para = doc.add_paragraph()
            target_level_para.add_run('Target Reading Level: ').bold = True
            target_level_para.add_run(f"Grade {target_level}")
            
            actual_level_para = doc.add_paragraph()
            actual_level_para.add_run('Actual Reading Level: ').bold = True
            actual_level_para.add_run(f"Grade {actual_level:.1f}")
            
            level_diff_para = doc.add_paragraph()
            level_diff_para.add_run('Difference: ').bold = True
            if abs(level_diff) <= 0.5:
                level_diff_para.add_run("On target")
            elif level_diff > 0:
                level_diff_para.add_run(f"{level_diff:+.1f} grades above target")
            else:
                level_diff_para.add_run(f"{level_diff:+.1f} grades below target")
        else:
            no_target_level_para = doc.add_paragraph()
            no_target_level_para.add_run('No target reading level specified. ')
            no_target_level_para.add_run(f'Current reading level: Grade {actual_level:.1f}').bold = True
        
        # Keyword Analysis
        if self.keyword_analysis:
            doc.add_heading('SEO Keyword Analysis', level=2)
            
            for keyword, count in self.keyword_analysis.items():
                keyword_para = doc.add_paragraph()
                keyword_para.add_run(f'"{keyword}": ').bold = True
                keyword_para.add_run(f"{count} occurrences")
        
        # Medicare Compliance (if applicable) - Enhanced with location info
        if self.user_config.get('is_medicare_page') and self.medicare_checks:
            doc.add_heading('Medicare Compliance Check', level=2)
            
            for check in self.medicare_checks:
                check_para = doc.add_paragraph()
                check_para.add_run(f"{check['type'].replace('_', ' ').title()}: ").bold = True
                check_para.add_run(check['description'])
                
                # Add location information if available
                if 'location' in check:
                    location_para = doc.add_paragraph()
                    location_para.add_run('Location: ').bold = True
                    location_para.add_run(check['location'])
        
        # Document Structure Guide
        if self.user_config.get('is_medicare_page'):
            doc.add_heading('Document Structure Guide', level=2)
            
            guide_para = doc.add_paragraph()
            guide_para.add_run('For optimal Medicare compliance, use these text markers in your document:').bold = True
            
            markers_list = [
                'start_page_copy - Begin main content processing',
                'end_page_copy - End main content processing', 
                'start_disclaimer - Begin disclaimer section (for CMS code placement)',
                'end_disclaimer - End disclaimer section'
            ]
            
            for marker in markers_list:
                marker_para = doc.add_paragraph()
                marker_para.add_run(f"• {marker}")
        
        # Detailed Changes Table - With fixed styling
        if results['detailed_corrections']:
            doc.add_heading('Detailed Changes Applied', level=2)
            
            # Create table with basic styling only
            table = doc.add_table(rows=1, cols=3)
            
            # Try to apply a basic table style, fall back to no style if it fails
            try:
                table.style = 'Table Grid'
            except:
                try:
                    table.style = 'Light Grid'
                except:
                    pass  # Use default table style
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Category'
            header_cells[1].text = 'Original Text'
            header_cells[2].text = 'Corrected Text'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add correction rows (limit to first 50 for readability)
            for correction in results['detailed_corrections'][:50]:
                row_cells = table.add_row().cells
                row_cells[0].text = correction['rule'].replace('_', ' ').title()
                row_cells[1].text = correction['original'][:100] + ('...' if len(correction['original']) > 100 else '')
                row_cells[2].text = correction['replacement'][:100] + ('...' if len(correction['replacement']) > 100 else '')
            
            if len(results['detailed_corrections']) > 50:
                more_para = doc.add_paragraph()
                more_para.add_run(f"... and {len(results['detailed_corrections']) - 50} more corrections")
                more_para.italic = True
        
        # Corrections by Category
        if results['corrections_by_category']:
            doc.add_heading('Corrections by Category', level=2)
            
            for category, count in sorted(results['corrections_by_category'].items(), 
                                        key=lambda x: x[1], reverse=True):
                if count > 0:
                    cat_para = doc.add_paragraph()
                    cat_para.add_run(f"{category.replace('_', ' ').title()}: ").bold = True
                    cat_para.add_run(f"{count} corrections")
    
    def apply_corporate_rules(self, doc):
        """Apply all corporate rules to document with formatting preservation"""
        total_corrections = 0
        corrections_by_category = defaultdict(int)
        detailed_corrections = []

        print(f"🔄 Applying corporate rules with formatting preservation...")

        # Enable Medicare rules if needed
        if self.user_config.get('is_medicare_page'):
            for rule in self.rules.get('medicare_rules', []):
                rule['enabled'] = True

        # Find bookmark range
        start_para, end_para = self._find_bookmark_range(doc)
        
        if start_para is None or end_para is None:
            print("📍 Bookmarks not found, processing entire document")
            paragraphs_to_process = doc.paragraphs
        else:
            print(f"📍 Processing content between bookmarks (paragraphs {start_para}-{end_para})")
            paragraphs_to_process = doc.paragraphs[start_para:end_para + 1]

        # Process each paragraph while preserving formatting
        for para_idx, paragraph in enumerate(paragraphs_to_process):
            if not paragraph.text.strip():
                continue

            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                
                original_text = run.text
                current_text = original_text
                
                # Extract and preserve angle bracket content
                current_text, bracket_placeholders = self._extract_angle_bracket_content(current_text)

                # Apply each rule category
                for category_name, category_rules in self.rules.items():
                    if not isinstance(category_rules, list):
                        continue

                    for rule in category_rules:
                        if not isinstance(rule, dict) or not rule.get('enabled', True):
                            continue

                        find_pattern = rule.get('find', '')
                        replacement = rule.get('replace', '')
                        is_function = rule.get('is_function', False)

                        if not find_pattern:
                            continue

                        try:
                            if is_function and callable(replacement):
                                new_text = re.sub(find_pattern, replacement, current_text, flags=re.IGNORECASE)
                            else:
                                case_sensitive = rule.get('case_sensitive', False)
                                flags = 0 if case_sensitive else re.IGNORECASE
                                new_text = re.sub(find_pattern, replacement, current_text, flags=flags)
                            
                            if new_text != current_text:
                                detailed_corrections.append({
                                    'category': category_name,
                                    'rule': rule.get('category', 'Unknown'),
                                    'original': current_text,
                                    'replacement': new_text
                                })
                                corrections_by_category[category_name] += 1
                                current_text = new_text
                                total_corrections += 1
                                
                        except re.error as e:
                            print(f"⚠️ Regex error in rule {rule.get('category', 'unknown')}: {e}")
                            continue

                # Restore angle bracket content
                current_text = self._restore_angle_bracket_content(current_text, bracket_placeholders)

                # Update run text if changes were made
                if current_text != original_text:
                    run.text = current_text

        # Apply post-processing for placeholders
        for paragraph in paragraphs_to_process:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue

                text = run.text
                original_text = text

                # Convert NUMBER_WORD_X to actual words
                if 'NUMBER_WORD_' in text:
                    number_words = {
                        'NUMBER_WORD_1': 'one', 'NUMBER_WORD_2': 'two', 'NUMBER_WORD_3': 'three',
                        'NUMBER_WORD_4': 'four', 'NUMBER_WORD_5': 'five', 'NUMBER_WORD_6': 'six',
                        'NUMBER_WORD_7': 'seven', 'NUMBER_WORD_8': 'eight', 'NUMBER_WORD_9': 'nine'
                    }
                    for placeholder, word in number_words.items():
                        text = text.replace(placeholder, word)

                if text != original_text:
                    run.text = text
                    total_corrections += 1

        return {
            'total_corrections': total_corrections,
            'corrections_by_category': dict(corrections_by_category),
            'detailed_corrections': detailed_corrections
        }
    
    def calculate_document_stats(self, doc):
        """Calculate document statistics"""
        start_para, end_para = self._find_bookmark_range(doc)
        
        if start_para is None or end_para is None:
            paragraphs_to_analyze = doc.paragraphs
        else:
            paragraphs_to_analyze = doc.paragraphs[start_para:end_para + 1]
        
        full_text = ' '.join([p.text for p in paragraphs_to_analyze if p.text.strip()])

        word_count = len(full_text.split())
        sentence_count = len(re.split(r'[.!?]+', full_text))
        paragraph_count = len([p for p in paragraphs_to_analyze if p.text.strip()])

        try:
            reading_level = textstat.flesch_kincaid_grade(full_text)
        except:
            reading_level = 0

        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'reading_level': reading_level,
            'full_text': full_text
        }
    
    def process_document(self, input_file_path, output_file_path, user_config):
        """Process document with user configuration"""
        try:
            self.user_config = user_config
            
            # Load document
            doc = Document(input_file_path)
            print(f"✅ Loaded document: {input_file_path}")
            
            # Calculate initial statistics
            stats = self.calculate_document_stats(doc)
            
            # Analyze keywords
            keywords = user_config.get('keywords', [])
            if keywords:
                self.keyword_analysis = self._analyze_keywords(stats['full_text'], keywords)
            
            # Check Medicare compliance
            if user_config.get('is_medicare_page'):
                self.medicare_checks = self._check_medicare_compliance(doc)
            
            # Apply corporate rules
            correction_results = self.apply_corporate_rules(doc)
            print(f"✅ Applied {correction_results['total_corrections']} corrections")
            
            # Recalculate statistics after processing
            final_stats = self.calculate_document_stats(doc)
            
            # Create comprehensive results
            results = {
                'success': True,
                'total_corrections': correction_results['total_corrections'],
                'corrections_by_category': correction_results['corrections_by_category'],
                'detailed_corrections': correction_results['detailed_corrections'],
                'document_statistics': final_stats,
                'keyword_analysis': self.keyword_analysis,
                'medicare_checks': self.medicare_checks
            }
            
            # Add analysis report to document
            self._create_analysis_report(doc, results)
            
            # Save processed document
            doc.save(output_file_path)
            print(f"✅ Saved processed document with analysis report: {output_file_path}")
            
            return results
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return {
                'success': False,
                'error': str(e)
            }

# Initialize processor
processor = MVPDocumentProcessor()

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_document():
    """Process uploaded document with user configuration"""
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload a .docx file'}), 400
        
        # Get user configuration
        user_config = {
            'target_word_count': request.form.get('target_word_count'),
            'keywords': [k.strip() for k in request.form.get('keywords', '').split(',') if k.strip()][:5],
            'target_reading_level': request.form.get('target_reading_level'),
            'is_medicare_page': request.form.get('is_medicare_page') == 'true'
        }
        
        # Convert numeric fields
        if user_config['target_word_count']:
            try:
                user_config['target_word_count'] = int(user_config['target_word_count'])
            except:
                user_config['target_word_count'] = None
        
        if user_config['target_reading_level']:
            try:
                user_config['target_reading_level'] = float(user_config['target_reading_level'])
            except:
                user_config['target_reading_level'] = None
        
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as input_temp:
            file.save(input_temp.name)
            input_path = input_temp.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as output_temp:
            output_path = output_temp.name
        
        # Process document
        results = processor.process_document(input_path, output_path, user_config)
        
        # Clean up input file
        os.unlink(input_path)
        
        if not results['success']:
            os.unlink(output_path)
            return jsonify({'error': f"Processing failed: {results['error']}"}), 500
        
        # Generate filename
        original_filename = secure_filename(file.filename)
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        processed_filename = f"{original_filename.rsplit('.', 1)[0]}_processed_{timestamp}.docx"
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=processed_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error in process_document: {e}")
        print(traceback.format_exc())
        return jsonify({'error': 'An unexpected error occurred during processing'}), 500

@app.route('/analyze', methods=['POST'])
def analyze_document():
    """Analyze document without processing (for preview)"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '' or not file or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file'}), 400
        
        # Get user configuration for analysis
        user_config = {
            'target_word_count': request.form.get('target_word_count'),
            'keywords': [k.strip() for k in request.form.get('keywords', '').split(',') if k.strip()][:5],
            'target_reading_level': request.form.get('target_reading_level'),
            'is_medicare_page': request.form.get('is_medicare_page') == 'true'
        }
        
        # Convert numeric fields
        if user_config['target_word_count']:
            try:
                user_config['target_word_count'] = int(user_config['target_word_count'])
            except:
                user_config['target_word_count'] = None
        
        if user_config['target_reading_level']:
            try:
                user_config['target_reading_level'] = float(user_config['target_reading_level'])
            except:
                user_config['target_reading_level'] = None
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
            file.save(temp.name)
            temp_path = temp.name
        
        try:
            # Load document for analysis
            doc = Document(temp_path)
            processor.user_config = user_config
            
            # Get statistics
            stats = processor.calculate_document_stats(doc)
            
            # Analyze keywords
            keyword_analysis = {}
            if user_config.get('keywords'):
                keyword_analysis = processor._analyze_keywords(stats['full_text'], user_config['keywords'])
            
            # Check Medicare compliance
            medicare_checks = []
            if user_config.get('is_medicare_page'):
                medicare_checks = processor._check_medicare_compliance(doc)
            
            # Count potential corrections (dry run)
            correction_preview = processor.apply_corporate_rules(doc)
            
            os.unlink(temp_path)
            
            return jsonify({
                'success': True,
                'document_statistics': stats,
                'potential_corrections': correction_preview['total_corrections'],
                'corrections_preview': correction_preview['corrections_by_category'],
                'keyword_analysis': keyword_analysis,
                'medicare_checks': medicare_checks,
                'user_config': user_config
            })
            
        except Exception as e:
            os.unlink(temp_path)
            raise e
            
    except Exception as e:
        print(f"Error in analyze_document: {e}")
        return jsonify({'error': 'Analysis failed'}), 500

@app.route('/health')
def health_check():
    """Health check endpoint for deployment"""
    return jsonify({
        'status': 'healthy',
        'service': 'MVP Document Processor Enhanced',
        'timestamp': datetime.datetime.now().isoformat()
    })

if __name__ == '__main__':
    # For local development
    app.run(debug=True, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
                